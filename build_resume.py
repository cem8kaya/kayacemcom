"""Build Cem Kaya executive resume — professional ATS 2-page .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0F, 0x17, 0x2A)
BLUE  = RGBColor(0x1D, 0x4E, 0xD8)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x47, 0x55, 0x69)
RULE  = RGBColor(0x93, 0xC5, 0xFD)

# ── Low-level helpers ─────────────────────────────────────────────────────────

def _font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def _sp(p, before=0, after=0, line=None):
    f = p.paragraph_format
    f.space_before = Pt(before)
    f.space_after  = Pt(after)
    if line:
        f.line_spacing = Pt(line)

def _rule(doc, color=RULE, weight="4", style="single"):
    p = doc.add_paragraph()
    _sp(p, before=0, after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   style)
    bot.set(qn('w:sz'),    weight)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*color))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def section(doc, title):
    _rule(doc)
    p = doc.add_paragraph()
    _sp(p, before=5, after=2)
    r = p.add_run(title.upper())
    _font(r, size=8.5, bold=True, color=BLUE)
    return p

def job(doc, title, company, period, location):
    p = doc.add_paragraph()
    _sp(p, before=6, after=0)
    r1 = p.add_run(title)
    _font(r1, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(f"  ·  {company}")
    _font(r2, size=10.5, color=BLUE)
    pm = doc.add_paragraph()
    _sp(pm, before=0, after=2)
    rm = pm.add_run(f"{period}  ·  {location}")
    _font(rm, size=8.5, italic=False, color=MUTED)

def bul(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    _sp(p, before=1, after=1)
    r = p.add_run(text)
    _font(r, size=9.5, color=DARK)

def skill_row(doc, cat, items):
    p = doc.add_paragraph()
    _sp(p, before=3, after=0)
    rc = p.add_run(cat + ":  ")
    _font(rc, size=9.5, bold=True, color=NAVY)
    ri = p.add_run("  ·  ".join(items))
    _font(ri, size=9.5, color=DARK)

# ── Document & margins ────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)
    sec.top_margin    = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin   = Cm(1.9)
    sec.right_margin  = Cm(1.9)

doc.styles['Normal'].font.name = "Calibri"
doc.styles['Normal'].font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
#  HEADER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
_sp(p, before=0, after=1)
r = p.add_run("CEM KAYA")
_font(r, size=26, bold=True, color=NAVY)

p = doc.add_paragraph()
_sp(p, before=0, after=3)
r = p.add_run("Telecom & AI Engineer  ·  5G / LTE / VoLTE  ·  AI/ML Systems  ·  Protocol Engineering")
_font(r, size=10.5, color=BLUE)

p = doc.add_paragraph()
_sp(p, before=0, after=1)
r = p.add_run("tcckaya8@gmail.com   |   linkedin.com/in/cemkaya   |   github.com/cem8kaya   |   Istanbul, Turkey")
_font(r, size=9, color=MUTED)

_rule(doc, color=NAVY, weight="8")

# ════════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
section(doc, "Executive Summary")
p = doc.add_paragraph()
_sp(p, before=2, after=5, line=13)
r = p.add_run(
    "Telecom & AI Engineer with 15+ years of experience delivering mission-critical network solutions "
    "and AI/ML systems for Tier-1 global operators. Proven record of leading complex end-to-end technical "
    "engagements — from 5G/LTE/VoLTE deep-dive root cause analysis and C-level performance reporting to "
    "the architecture of production AI platforms — across customers including Deutsche Telekom, AT&T, "
    "Verizon, Telia Sonera, O2, Ooredoo, Vodafone, and Turkcell. At B-YOND, served as primary technical "
    "resource on multi-operator analytics programmes and contributed directly to C-level strategic "
    "reporting and R&D initiatives. At Turkcell, conceived and delivered internal automation and "
    "intelligence systems that transformed executive network operations decisions. Currently at Odine Labs, "
    "architecting next-generation AI systems — multi-agent LLM platforms, synthetic data pipelines, and "
    "V2X intelligence — at the frontier of telecom and AI convergence."
)
_font(r, size=9.8, color=DARK)

# ════════════════════════════════════════════════════════════════════════════
#  EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════
section(doc, "Professional Experience")

# ── Odine Labs ──
job(doc,
    "Solution Architect",
    "Odine Labs R&D",
    "April 2025 – Present",
    "Istanbul, Turkey")
for b in [
    "Architecting a cloud-native AI/ML platform on Kubernetes/OpenShift and GCP Vertex AI, delivering "
    "40% faster model deployment cycles for telecom R&D use cases.",
    "Designed and led delivery of a conversational AI assistant and intelligent document analyzer using "
    "LangGraph multi-agent orchestration and FAISS-based RAG pipelines over 3GPP specification corpora.",
    "Built a statistically rigorous synthetic telecom data generation engine (Gamma/Lognormal/Beta/Poisson, "
    "ARIMA time series, five anomaly modes) underpinning ML training and validation workflows.",
    "Developed V2X sidelink QoS prediction models (XGBoost / Neural Networks) for connected mobility "
    "research, advancing the company's 5G-Advanced and C-V2X product portfolio.",
]:
    bul(doc, b)

# ── B-YOND ──
job(doc,
    "Senior Data Analyst & Scientist",
    "B-YOND",
    "December 2022 – March 2025",
    "Canada (Remote)")
for b in [
    "Led end-to-end deep-dive analytics and root cause analysis engagements for Tier-1 operators "
    "including Deutsche Telekom, AT&T, Verizon, Telia Sonera, O2, and Ooredoo, covering 5G SA/NSA, "
    "LTE, and VoLTE network domains.",
    "Developed 5G Core and VoNR mobility & connectivity prediction models for Verizon, improving "
    "reliability KPIs by 25%; findings fed directly into quarterly C-level performance reviews.",
    "Architected EPC and 5GC predictive models for AT&T's 4G-to-5G core migration, enabling "
    "proactive capacity management across a multi-phase, multi-region network transition.",
    "Delivered a real-time anomaly detection system sustaining millions of concurrent sessions during "
    "the FIFA World Cup 2023 (Ooredoo Qatar) — zero critical core incidents throughout the tournament.",
    "Built automated VoLTE root cause diagnosis models for Deutsche Telekom group operators (Telenet, "
    "Vodafone Ziggo), substantially reducing mean time to resolution for voice quality degradations.",
    "Contributed to R&D roadmap and internal C-level technical reporting frameworks; developed "
    "multi-protocol parsers (SIP, Diameter, GTPv2, HTTP/2, PFCP, NGAP) as platform core components.",
]:
    bul(doc, b)

# ── Turkcell ──
job(doc,
    "Senior Core Network Optimization Engineer",
    "Turkcell",
    "September 2019 – February 2022",
    "Istanbul, Turkey")
for b in [
    "Conceived and built from scratch a network operations automation and ML-driven anomaly detection "
    "platform (Python, SQL, JavaScript) reducing manual core network interventions by 70%; outcomes "
    "reported at executive and C-level operational reviews.",
    "Implemented Meta Prophet time-series models for predictive 5G/LTE performance forecasting, "
    "feeding proactive capacity planning and C-level KPI dashboards.",
    "Spearheaded VoiceX — a strategic multi-vendor IMS migration (Ericsson/Huawei → Mavenir "
    "cloud-native) executed with zero service downtime across millions of VoLTE subscribers.",
    "Unified Ericsson ENM, Mavenir XA, and Huawei OSS/BSS platforms into a single operational "
    "intelligence layer, consolidating visibility across the full 5G/LTE/IMS stack.",
    "Led internal R&D investigations into 5G core anomaly signatures and automated root cause "
    "classification as part of Turkcell's AI-for-Networks strategic programme.",
]:
    bul(doc, b)

# ── Early Career ──
job(doc,
    "Core Network & VoIP Engineer",
    "Vodafone Turkey  ·  VOTEL  ·  Self-Employed  ·  Turkcell",
    "2010 – 2019",
    "Turkey / Various")
for b in [
    "Managed mission-critical IMS, PSTN gateway, and SBC infrastructure serving tens of millions of "
    "subscribers across multiple national operators.",
    "Ran an independent VoIP and telecommunications consulting practice; delivered end-to-end solutions "
    "on Oracle ACME SBC, GENBAND, and Asterisk platforms for carrier and enterprise clients.",
]:
    bul(doc, b)

# ════════════════════════════════════════════════════════════════════════════
#  TECHNICAL SKILLS
# ════════════════════════════════════════════════════════════════════════════
section(doc, "Technical Skills")

skill_row(doc, "5G / Telecom Protocols",
    ["5G SA/NSA", "LTE/EPC", "VoLTE/VoNR", "IMS", "AMF/SMF/UPF/PCF/UDM/AUSF/NRF/NWDAF",
     "SIP", "DIAMETER", "GTPv1/v2", "PFCP", "NGAP/S1AP/NAS", "RTP/RTCP", "HTTP/2 (SBA)",
     "TS 23.501/502/503", "TS 23.288", "TS 29.244", "TS 33.501/513"])
skill_row(doc, "AI / ML",
    ["LangGraph Multi-Agent", "RAG / FAISS", "QLoRA Fine-Tuning", "HuggingFace PEFT",
     "Isolation Forest", "DTW Pattern Matching", "Anomaly Detection", "Meta Prophet",
     "XGBoost", "Time-Series Forecasting"])
skill_row(doc, "Languages",
    ["Python", "C++ (C++17)", "Swift", "JavaScript (ES6+)", "Bash", "SQL", "R"])
skill_row(doc, "Frameworks & Tools",
    ["Open5GS", "UERANSIM", "Flask", "FastAPI", "Scapy", "Docker", "Kubernetes",
     "CMake", "libpcap", "nDPI", "OpenSSL"])
skill_row(doc, "Cloud & Infra",
    ["GCP", "GCP Vertex AI", "OpenStack", "Prometheus", "Grafana",
     "MongoDB", "Elasticsearch", "Jenkins / CI/CD"])
skill_row(doc, "Mobile",
    ["iOS / Swift", "CoreMotion", "AVFoundation", "On-Device ML", "App Store"])

# ════════════════════════════════════════════════════════════════════════════
#  PORTFOLIO HIGHLIGHTS  (brief, to stay on 2 pages)
# ════════════════════════════════════════════════════════════════════════════
section(doc, "Open-Source Portfolio  ·  github.com/cem8kaya")

projects = [
    ("Open5GS AI/ML Laboratory",
     "Full 5G SA core (7 NFs) on GCP + UERANSIM. 8 labeled ML traffic scenarios, ETL pipeline, "
     "Isolation Forest / Random Forest."),
    ("Open5GS NWDAF — C++ Reference",
     "Standalone C++ NWDAF per 3GPP TS 23.288 v17. All 7 analytics IDs, embedded Isolation Forest + "
     "EWMA, full TS 29.520 SBI REST API, systemd service."),
    ("Digital Twin RCA Agent",
     "LangGraph multi-agent 5G RCA system. DTW matching across 15 KPIs, FAISS RAG over 3GPP specs, "
     "LLM failover chain. <3 min MTRCA."),
    ("Callflow Visualizer — Enhanced DPI",
     "C++17 PCAP analyser; 15+ protocol parsers (SIP/GTP/PFCP/NGAP/NAS/HTTP2), D3.js ladder "
     "diagrams, ~34,700 pkt/s."),
    ("5GC Secure CUPS Shield",
     "PFCP firewall per TS 29.244 / TS 33.513. HMAC-SHA256 auth, DoS protection, dynamic threat scoring."),
    ("5G LLM Engine",
     "QLoRA fine-tuned telecom LLM + FAISS RAG over 3GPP specs. FastAPI RCA output for AMF/SMF/IMS/GTP."),
]

for title, desc in projects:
    p = doc.add_paragraph()
    _sp(p, before=3, after=0)
    r1 = p.add_run(f"{title}  —  ")
    _font(r1, size=9.2, bold=True, color=NAVY)
    r2 = p.add_run(desc)
    _font(r2, size=9.2, color=DARK)

# ════════════════════════════════════════════════════════════════════════════
#  EDUCATION & CERTIFICATIONS
# ════════════════════════════════════════════════════════════════════════════
section(doc, "Education")
for deg, school, loc in [
    ("M.Sc. Management Information Systems", "Bilgi University", "Istanbul, Turkey"),
    ("B.Sc. Electronics and Communication Engineering", "Kocaeli University", "Turkey"),
]:
    p = doc.add_paragraph()
    _sp(p, before=3, after=0)
    r1 = p.add_run(f"{deg}  ·  ")
    _font(r1, size=9.5, bold=True, color=NAVY)
    r2 = p.add_run(f"{school}, {loc}")
    _font(r2, size=9.5, color=DARK)

section(doc, "Certifications")
p = doc.add_paragraph()
_sp(p, before=2, after=2)
r = p.add_run(
    "IELTS 7.0  ·  Nokia Vo5G Protocol Stack & Standards / IMS Protocols  ·  "
    "3GPP 5G Security Specification  ·  Udacity Business Analyst  ·  Kaggle ML & Python"
)
_font(r, size=9, color=DARK)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/kayacemcom/origresumes/Cem_Kaya_Resume_2026.docx"
doc.save(out)
print(f"Saved → {out}")
